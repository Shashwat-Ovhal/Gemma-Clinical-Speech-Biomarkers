"""
Patches research_paper_ieee_jbhi.docx with 4 changes:
  1. New paper title (everywhere in the doc)
  2. New abstract text
  3. Text Box 1 reformatted as a proper styled block
  4. GitHub URL in place of placeholder
"""
import copy
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re

DOCX_PATH = r"manuscript/research_paper_ieee_jbhi.docx"

OLD_TITLE = "MedGemma-PD: An Edge-Distilled, XAI-Driven Acoustic Pipeline for Robust Unconstrained Parkinson\u2019s Disease Telemetry"
NEW_TITLE = "Wavelet-Anchored Cross-Attention Fusion with Edge Knowledge Distillation for Robust Parkinson\u2019s Disease Acoustic Biomarker Analysis"

NEW_ABSTRACT = (
    "Background and Objective: Continuous remote monitoring of Parkinson\u2019s Disease (PD) via mobile health (mHealth) "
    "telemetry remains clinically impractical due to unconstrained acoustic noise, demographic confounding, and a critical "
    "gap between raw model predictions and actionable clinical output. This paper presents MedGemma-PD, an end-to-end, "
    "edge-deployable acoustic pipeline engineered for robust PD detection from uncontrolled smartphone recordings.\n"
    "Methods: Using the mPower longitudinal cohort (N=120, post age-gating \u226545 years), we applied a Tunable Q-Factor "
    "Wavelet Transform (TQWT; Q=3, r=3, J=12) to isolate laryngeal perturbation features from ambient noise. A novel "
    "Cross-Attention Fusion module anchors frozen wav2vec 2.0 embeddings to deterministic physiological features\u2014Jitter, "
    "Shimmer, HNR, and F0std\u2014to counter overfitting on small clinical datasets. The full model was compressed via knowledge "
    "distillation (T=4.0, \u03b1=0.7) into a 287K-parameter 1D-CNN edge student. An on-device, 4-bit quantized Gemma-2b agent "
    "synthesizes SHAP-driven outputs into structured clinical narratives. Evaluation employed 5-fold subject-stratified "
    "GroupKFold cross-validation to prevent data leakage.\n"
    "Results: The TQWT-denoised physiological pipeline achieved a peak AUC-ROC of 0.700 [95% CI: 0.606\u20130.793], significantly "
    "outperforming full neural fine-tuning (AUC: 0.502) and MFCC baselines (AUC: 0.332), both of which collapsed under "
    "real-world noise. The edge student achieved 31.9 \u00b1 25.5 ms CPU inference latency with a 1.1 MB footprint, representing "
    "a >370\u00d7 compression over the teacher model. LLM-generated clinical narratives scored 4.8/5.0 for factual accuracy with "
    "excellent inter-rater agreement (Cohen\u2019s \u03ba = 0.88).\n"
    "Conclusions: MedGemma-PD demonstrates that physiologically-grounded, noise-resistant feature engineering outperforms deep "
    "learning in unconstrained mHealth settings. The framework delivers an edge-deployable, HIPAA-aligned pipeline that bridges "
    "acoustic telemetry and clinical decision support for scalable PD telemonitoring."
)

TEXTBOX_LINES = [
    ("[Assessment]: At Risk (Risk Signal: 0.78)", True),
    ("Analysis of speech biomarkers indicates elevated motor control risk.", False),
    ("", False),
    ("[Evidence Integration]:", True),
    ("1. Speech Biomarkers: Jitter: 0.021; Shimmer: 0.052; HNR: 15.4 dB", False),
    ("2. Longitudinal Context: UPDRS Trend: Deteriorating (Change: +3.5)", False),
    ("3. Key XAI Drivers: Model attention focused on sustained phonation (vowel /a/) driven by elevated amplitude perturbation (Shimmer).", False),
    ("", False),
    ("[Recommendation]: Schedule Neurology Review within 14 days.", True),
]

GITHUB_URL = "https://github.com/Shashwat-Ovhal/Gemma-Clinical-Speech-Biomarkers"

doc = Document(DOCX_PATH)

# ── Helper: replace text inside a paragraph preserving runs/formatting ──────
def replace_para_text(para, new_text):
    """Replace all text in a paragraph while preserving the first run's formatting."""
    # Grab format from the first run
    if para.runs:
        fmt_run = para.runs[0]
        bold = fmt_run.bold
        italic = fmt_run.italic
        font_name = fmt_run.font.name
        font_size = fmt_run.font.size
    else:
        bold = italic = False
        font_name = font_size = None

    # Clear all existing runs
    for run in para.runs:
        run.text = ""

    # Set entire text in first run (or add one)
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run()

    r.text = new_text
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if font_name:
        r.font.name = font_name
    if font_size:
        r.font.size = font_size


# ── 1. Replace title everywhere it appears ──────────────────────────────────
print("1) Replacing title...")
for para in doc.paragraphs:
    if OLD_TITLE in para.text or "MedGemma-PD: An Edge-Distilled" in para.text:
        # Preserve existing run formatting
        for run in para.runs:
            if OLD_TITLE in run.text or "MedGemma-PD: An Edge-Distilled" in run.text:
                run.text = run.text.replace(
                    "MedGemma-PD: An Edge-Distilled, XAI-Driven Acoustic Pipeline for Robust Unconstrained Parkinson\u2019s Disease Telemetry",
                    NEW_TITLE
                ).replace(
                    "MedGemma-PD: An Edge-Distilled, XAI-Driven Acoustic Pipeline for Robust Unconstrained Parkinson's Disease Telemetry",
                    NEW_TITLE
                )
        # Also handle single-run paragraphs where text is split across runs oddly
        if "MedGemma-PD: An Edge-Distilled" in para.text:
            replace_para_text(para, NEW_TITLE)

# Also patch in tables (if title appears in any table cell)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if "MedGemma-PD: An Edge-Distilled" in run.text:
                        run.text = run.text.replace(
                            "MedGemma-PD: An Edge-Distilled, XAI-Driven Acoustic Pipeline for Robust Unconstrained Parkinson\u2019s Disease Telemetry",
                            NEW_TITLE
                        )


# ── 2. Replace abstract ──────────────────────────────────────────────────────
print("2) Replacing abstract...")
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("Background and Objective:") and "mHealth" in para.text:
        replace_para_text(para, NEW_ABSTRACT)
        break


# ── 3. Reformat Text Box 1 ──────────────────────────────────────────────────
print("3) Reformatting Text Box 1...")
textbox_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Text Box 1. Autonomous Agentic Output Example" in para.text:
        textbox_idx = i
        break

if textbox_idx is not None:
    tb_para = doc.paragraphs[textbox_idx]

    # Build new text: label + code block content cleanly (no markdown >)
    label_text = "Text Box 1. Autonomous Agentic Output Example"
    box_content = (
        "[Assessment]: At Risk (Risk Signal: 0.78)\n"
        "Analysis of speech biomarkers indicates elevated motor control risk.\n\n"
        "[Evidence Integration]:\n"
        "  1. Speech Biomarkers: Jitter: 0.021; Shimmer: 0.052; HNR: 15.4 dB\n"
        "  2. Longitudinal Context: UPDRS Trend: Deteriorating (Change: +3.5)\n"
        "  3. Key XAI Drivers: Model attention focused on sustained phonation (vowel /a/)\n"
        "     driven by elevated amplitude perturbation (Shimmer).\n\n"
        "[Recommendation]: Schedule Neurology Review within 14 days."
    )

    # Clear existing content and write label (bold) + box content
    for run in tb_para.runs:
        run.text = ""

    if tb_para.runs:
        label_run = tb_para.runs[0]
    else:
        label_run = tb_para.add_run()

    label_run.text = label_text
    label_run.bold = True

    content_run = tb_para.add_run("\n" + box_content)
    content_run.bold = False
    content_run.font.name = "Courier New"
    content_run.font.size = Pt(9)


# ── 4. Replace GitHub placeholder ───────────────────────────────────────────
print("4) Replacing GitHub placeholder...")
placeholder_patterns = [
    "[PLACEHOLDER \u2014 Insert GitHub URL or Zenodo DOI here]",
    "[PLACEHOLDER — Insert GitHub URL or Zenodo DOI here]",
    "PLACEHOLDER",
]

for para in doc.paragraphs:
    if "PLACEHOLDER" in para.text and ("GitHub" in para.text or "Zenodo" in para.text or "codebase" in para.text.lower()):
        new_text = para.text
        for pat in placeholder_patterns:
            new_text = new_text.replace(pat, GITHUB_URL)
        replace_para_text(para, new_text)


# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"\nSaved: {DOCX_PATH}")
print("Done. All 4 changes applied.")
